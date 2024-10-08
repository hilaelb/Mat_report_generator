import os
import json
import re
from datetime import datetime
import time
import docx.opc.constants
import scipy.io
from pprint import pprint
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import numpy as np
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont, QIcon
from PyQt5.QtWidgets import QApplication, QDialog, QFileDialog, QVBoxLayout, QCheckBox, QPushButton, QLabel, QScrollArea, QWidget, QHBoxLayout,QListWidget, QProgressDialog, QMessageBox
from plots import create_plot, create_kml_from_waypoints, kml_to_gmplot_image, create_histogram, create_partitioned_plot, create_mixed_plot, plot_tiger_modes, create_3d_plot


def extract_datetime_from_path(plot_path):
    # Define a regex pattern to match the date and time in the format: YYYY_MM_DD-HH_MM_SS
    pattern = r'(\d{4}_\d{2}_\d{2})-(\d{2}_\d{2}_\d{2})'

    # Search the plot_path for the pattern
    match = re.search(pattern, plot_path)

    if match:
        # Extract the date and time strings
        date_str = match.group(1)
        time_str = match.group(2)

        # Convert to datetime object
        datetime_obj = datetime.strptime(f'{date_str} {time_str}', '%Y_%m_%d %H_%M_%S')

        # Format to desired output
        formatted_datetime = datetime_obj.strftime('%A, %m/%d/%y, %H:%M:%S')
        return formatted_datetime
    else:
        return None

def create_word_document(plot_paths, doc_file_path,map_file_paths):
    doc = Document()

    # Add the cover page
    doc.add_heading('Experimental Report of Tiger', level=1).alignment = 1  # Centered heading
    # Add the LAR.png image and center it
    lar_paragraph = doc.add_paragraph('\n \n \n')
    lar_run = lar_paragraph.add_run()
    lar_run.add_picture('LAR.png', width=Inches(2.5), height=Inches(2.5))
    lar_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break for the table of contents
    doc.add_page_break()

    last_file_date = None
    experiment_counter = 1

    # Define maximum dimensions for the images (in inches)
    max_width = Inches(6.5)  # Slightly less than the page width
    max_height = Inches(9)  # Slightly less than the page height

    for index, (plot_path, description) in enumerate(plot_paths):
        # Extract the filename to check for changes
        current_file_date = extract_datetime_from_path(plot_path)

        if current_file_date != last_file_date and current_file_date is not None:
            if index > 0:
                doc.add_page_break()
            doc.add_heading(f'Experiment {experiment_counter}:', level=1)
            doc.add_paragraph('Location:')
            doc.add_paragraph(f'Date and time: {current_file_date}')
            doc.add_paragraph('Comments:')
            last_file_date = current_file_date
            experiment_counter += 1

        try:
            doc.add_picture(plot_path, width=Inches(6))
            if description != '':
                doc.add_paragraph(description).alignment = 1  # Centered text

            # Check if there's a corresponding map file for this plot
            base_name = os.path.basename(plot_path).replace('.png', '')
            map_file_path = next((path for path in map_file_paths if base_name in os.path.basename(path)), None)

            if map_file_path:
                doc.add_paragraph('Working Area:')
                doc.add_picture(map_file_path, width=Inches(6))


        except Exception as e:
            print(f"Error adding plot {plot_path} to the document: {e}")
            continue


    doc.save(doc_file_path)

def clear_folder(plots_folder):
    for file_name in os.listdir(plots_folder):
        file_path = os.path.join(plots_folder, file_name)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')


def generate_report(selected_plots, file_paths):

    try:
        desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        # output_folder = os.path.join(desktop_path, 'MAT_Files_Report')

        app = QApplication([])  # Create a PyQt application instance

        # Show a dialog to select the output folder
        output_folder = QFileDialog.getExistingDirectory(None, "Select Output Folder", os.path.expanduser("~"))

        if not output_folder:
            output_folder = os.path.join(desktop_path, 'MAT_Files_Report')


        # Show a dialog to specify the name of the Word document
        word_file_name, _ = QFileDialog.getSaveFileName(None, "Save Word Document As",
                                                        os.path.join(output_folder, "Report.docx"),
                                                        "Word Documents (*.docx)")

        if not word_file_name:
            print("No file name provided. Report generation canceled.")
            return

        plots_folder = os.path.join(output_folder, 'plots')
        maps_folder = os.path.join(output_folder, 'maps')
        kmls_folder = os.path.join(output_folder, 'kmls')


        os.makedirs(maps_folder, exist_ok=True)
        os.makedirs(plots_folder, exist_ok=True)
        # os.makedirs(output_folder, exist_ok=True)
        os.makedirs(kmls_folder, exist_ok=True)


        clear_folder(plots_folder)
        clear_folder(maps_folder)
        clear_folder(kmls_folder)


        # Create and show a progress dialog
        progress_dialog = QProgressDialog("Generating report...", "Cancel", 0, len(file_paths), None)
        progress_dialog.setGeometry(300, 300, 500, 200)
        progress_dialog.setWindowTitle("Report Generation")
        progress_dialog.setWindowModality(Qt.WindowModal)
        progress_dialog.setMinimumDuration(0)
        progress_dialog.setValue(0)
        progress_dialog.show()


        try:

            plot_paths = []
            map_paths = []
            kmls = []

            for idx, mat_file_path in enumerate(file_paths):
                base_name = os.path.basename(mat_file_path).replace('.mat', '')

                # Update progress dialog with current file processing status
                progress_dialog.setLabelText(f"Processing file {base_name}...")
                QApplication.processEvents()

                try:

                    if selected_plots[mat_file_path]["Depth/Time Plot"]:
                        plot_file_name_1 = f'{base_name}_ba30_depth_vs_time.png'
                        plot_file_path_1 = os.path.join(plots_folder, plot_file_name_1)
                        dict = {'ba30': ['depth', 'timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_plot(plot_file_path_1, dict_value_items, title='Depth over time')
                        plot_paths.append((plot_file_path_1, ''))

                        # Update progress dialog with the specific plot status
                        progress_dialog.setLabelText(f"Building Depth/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Lat/Long Plot"]:
                        plot_file_name_2 = f'{base_name}_waypoint_rosbag_lat_vs_long.png'
                        plot_file_path_2 = os.path.join(plots_folder, plot_file_name_2)
                        dict = {'waypoint': ['latitude', 'longitude']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        pprint(dict_value_items)
                        create_plot(plot_file_path_2, dict_value_items, plot=True)
                        plot_paths.append((plot_file_path_2, ''))

                        progress_dialog.setLabelText(f"Building Lat/Long Plot for {base_name}...")
                        QApplication.processEvents()

                        kml_path = os.path.join(kmls_folder, f'{base_name}_waypoint_rosbag_lat_vs_long.kml')
                        create_kml_from_waypoints(mat_file_path, 'waypoint', 'latitude', 'longitude', kml_path)
                        kmls.append(kml_path)

                        progress_dialog.setLabelText(f"Building KML for {base_name}...")
                        QApplication.processEvents()

                        if selected_plots[mat_file_path]["Add a picture of the map"]:
                            map_file_path = os.path.join(maps_folder, f'{base_name}_waypoint_rosbag_lat_vs_long.png')
                            kml_to_gmplot_image(mat_file_path, 'waypoint', 'latitude', 'longitude', map_file_path)
                            map_paths.append(map_file_path)

                            progress_dialog.setLabelText(f"Building Map for {base_name}...")
                            QApplication.processEvents()


                    if selected_plots[mat_file_path]["Jet Rpm/Time Plot"]:
                        plot_file_name_5 = f'{base_name}_jet_rpm_vs_time.png'
                        plot_file_path_5 = os.path.join(plots_folder, plot_file_name_5)
                        dict = {'jet': ['rpm', 'timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_plot(plot_file_path_5, dict_value_items, 'Jet RPM')
                        plot_paths.append((plot_file_path_5, ''))

                        progress_dialog.setLabelText(f"Building Jet Rpm/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Jet Voltage/Time Plot"]:
                        plot_file_name_6 = f'{base_name}_jet_voltage_vs_time.png'
                        plot_file_path_6 = os.path.join(plots_folder, plot_file_name_6)
                        dict = {'jet': ['v_in', 'timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_plot(plot_file_path_6, dict_value_items, 'Jet Voltage', range=(50, 60))
                        plot_paths.append((plot_file_path_6, ''))

                        progress_dialog.setLabelText(f"Building Jet Voltage/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Jet data"]:
                        plot_file_name_7 = f'{base_name}_jet_data.png'
                        plot_file_path_7 = os.path.join(plots_folder, plot_file_name_7)
                        dict = {'jet': ['timestamp', 'avg_input_current', 'avg_motor_current', 'rpm', 'temp_fet', 'v_in']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_7, dict_value_items, 'Jet data over time')
                        plot_paths.append((plot_file_path_7, ''))

                        progress_dialog.setLabelText(f"Building Jet Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Servo Angle/Time Plot"]:
                        plot_file_name_8 = f'{base_name}_servo_angle_vs_time.png'
                        plot_file_path_8 = os.path.join(plots_folder, plot_file_name_8)
                        dict = {'servo': ['timestamp', 'srv1_angle', 'srv2_angle', 'srv3_angle', 'srv4_angle']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_8, dict_value_items, 'Servos angle over time')
                        plot_paths.append((plot_file_path_8, ''))

                        progress_dialog.setLabelText(f"Building Servo Angle/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Servo Voltage/Time Plot"]:
                        plot_file_name_9 = f'{base_name}_servo_voltage_vs_time.png'
                        plot_file_path_9 = os.path.join(plots_folder, plot_file_name_9)
                        dict = {'servo': ['timestamp', 'srv1_voltage', 'srv2_voltage', 'srv3_voltage', 'srv4_voltage']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_9, dict_value_items, 'Servos Voltage over time')
                        plot_paths.append((plot_file_path_9, ''))

                        progress_dialog.setLabelText(f"Building Servo Voltage/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Servo Temperature/Time Plot"]:
                        plot_file_name_10 = f'{base_name}_servo_temperature_vs_time.png'
                        plot_file_path_10 = os.path.join(plots_folder, plot_file_name_10)
                        dict = {'servo': ['timestamp', 'srv1_temperature', 'srv2_temperature', 'srv3_temperature',
                                          'srv4_temperature']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_10, dict_value_items, 'Servos temperature over time')
                        plot_paths.append((plot_file_path_10, ''))

                        progress_dialog.setLabelText(f"Building Servo Temperature/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Inertial data"]:
                        plot_file_name_11 = f'{base_name}_inertial_data.png'
                        plot_file_path_11 = os.path.join(plots_folder, plot_file_name_11)
                        dict = {'insstatus': ['timestamp', 'ins_mode', 'ins_error', 'ins_fix']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_mixed_plot(plot_file_path_11, dict_value_items, title='Inertial system data',
                                          y_names='Fix & Error & Mode')
                        plot_paths.append((plot_file_path_11, ''))

                        progress_dialog.setLabelText(f"Building Inertial Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Gyro/Accel data"]:
                        plot_file_name_12 = f'{base_name}_gyro_data.png'
                        plot_file_path_12 = os.path.join(plots_folder, plot_file_name_12)
                        dict = {'gyro_accel_data': ['timestamp', 'gyro_x', 'gyro_y', 'gyro_z', 'accel_x', 'accel_y', 'accel_z']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_12, dict_value_items, 'Gyro/Accel data over time')
                        plot_paths.append((plot_file_path_12, ''))

                        progress_dialog.setLabelText(f"Building Gyro/Accel Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["BME front measurements"]:
                        plot_file_name_14 = f'{base_name}_bme_front_data.png'
                        plot_file_path_14 = os.path.join(plots_folder, plot_file_name_14)
                        dict = {'bme_front': ['timestamp', 'temperature', 'pressure', 'humidity']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_14, dict_value_items, 'BME front measurements')
                        plot_paths.append((plot_file_path_14, ''))

                        progress_dialog.setLabelText(f"Building BME Front Data Plot for {base_name}...")
                        QApplication.processEvents()


                    if selected_plots[mat_file_path]["BME back measurements"]:
                        plot_file_name_15 = f'{base_name}_bme_back_data.png'
                        plot_file_path_15 = os.path.join(plots_folder, plot_file_name_15)
                        dict = {'bme_back': ['timestamp', 'pressure']}
                        dict2 = {'bme_back': ['temperature', 'humidity']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        dict_value_items2 = extract_values_from_data(mat_file_path, dict2)
                        create_3d_plot(plot_file_path_15, dict_value_items, dict_value_items2, 'BME back measurements')
                        # create_mixed_plot(plot_file_path_15,dict_value_items,title='BME back measurements',y_names='Temperature & Pressure & Humidity')
                        plot_paths.append((plot_file_path_15, ''))

                        progress_dialog.setLabelText(f"Building BME Back Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Bar30 measurements"]:
                        plot_file_name_16 = f'{base_name}_bar30_data.png'
                        plot_file_path_16 = os.path.join(plots_folder, plot_file_name_16)
                        dict = {'ba30': ['timestamp', 'temperature', 'pressure', 'depth']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_mixed_plot(plot_file_path_16, dict_value_items, title='Bar30 measurements',
                                          y_names='Temperature & Pressure & Depth')
                        # create_partitioned_plot(plot_file_path_16,dict_value_items,'Bar30 measurements')
                        plot_paths.append((plot_file_path_16, ''))

                        progress_dialog.setLabelText(f"Building Bar30 Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Motor Guidance data"]:
                        plot_file_name_17 = f'{base_name}_motor_guidance_data.png'
                        plot_file_path_17 = os.path.join(plots_folder, plot_file_name_17)
                        dict = {'motor_guidance_data': ['timestamp', 'srv_1_angle', 'srv_2_angle', 'srv_4_angle', 'srv_4_angle',
                                                        'jet_rpm']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_partitioned_plot(plot_file_path_17, dict_value_items, 'Motor guidance data - angels over time')
                        plot_paths.append((plot_file_path_17, ''))

                        progress_dialog.setLabelText(f"Building Motor Guidance Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Vn200 data"]:
                        plot_file_name_18 = f'{base_name}_vn200_data.png'
                        plot_file_path_18 = os.path.join(plots_folder, plot_file_name_18)
                        dict = {'vn': ['yaw', 'timestamp']}
                        dict2 = {'vn': ['timestamp', 'velned_x', 'velned_y', 'velned_z']}
                        dict3 = {'vn': ['timestamp', 'poslla_x', 'poslla_y', 'poslla_z']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        dict_value_items2 = extract_values_from_data(mat_file_path, dict2)
                        dict_value_items3 = extract_values_from_data(mat_file_path, dict3)
                        create_plot(plot_file_path_18, dict_value_items, 'Vn200 data - yaw over time', range=(-200, 200))
                        create_mixed_plot(f'{plot_file_path_18[:-4]}_2.png', dict_value_items2,
                                          title='Vn200 data - velocity over time', y_names='Velocity x & y & z')
                        create_mixed_plot(f'{plot_file_path_18[:-4]}_3.png', dict_value_items3,
                                          title='Vn200 data - position over time', y_names='position x & y & z', )
                        plot_paths.append((plot_file_path_18, ''))
                        plot_paths.append((f'{plot_file_path_18[:-4]}_2.png', ''))
                        plot_paths.append((f'{plot_file_path_18[:-4]}_3.png', ''))

                        progress_dialog.setLabelText(f"Building Vn200 Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Histogram of Acceleration"]:
                        plot_file_name_19 = f'{base_name}_acceleration_histogram.png'
                        plot_file_path_19 = os.path.join(plots_folder, plot_file_name_19)
                        dict = {'gyro_accel_data': ['accel_x', 'accel_y', 'accel_z']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_19)
                        plot_paths.append((plot_file_path_19, ''))

                        progress_dialog.setLabelText(f"Building Acceleration Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Tiger Mode/Time Plot"]:
                        plot_file_name_7 = f'{base_name}_tiger_mode_vs_time.png'
                        plot_file_path_7 = os.path.join(plots_folder, plot_file_name_7)
                        dict = {'tiger': ['tiger_mode', 'timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        plot_tiger_modes(dict_value_items, plot_file_path_7)
                        plot_paths.append((plot_file_path_7, ''))

                        progress_dialog.setLabelText(f"Building Tiger Mode/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["QGC Mode/Time Plot"]:
                        plot_file_name_20 = f'{base_name}_qgc_mode_vs_time.png'
                        plot_file_path_20 = os.path.join(plots_folder, plot_file_name_20)
                        dict = {'qgc': ['mode', 'timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        plot_tiger_modes(dict_value_items, plot_file_path_20)
                        plot_paths.append((plot_file_path_20, ''))

                        progress_dialog.setLabelText(f"Building QGC Mode/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["ADC measurements"]:
                        plot_file_name_13 = f'{base_name}_adc_data.png'
                        plot_file_path_13 = os.path.join(plots_folder, plot_file_name_13)
                        dict = {'adc': ['timestamp', 'v_24', 'v_12', 'v_5', 'v_3']}
                        dict2 = {'adc': ['timestamp', 'i_24', 'i_12', 'i_5', 'i_3']}
                        dict3 = {'adc': ['v_in', 'timestamp']}
                        dict4 = {'adc': ['i_in', 'timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        dict_value_items2 = extract_values_from_data(mat_file_path, dict2)
                        dict_value_items3 = extract_values_from_data(mat_file_path, dict3)
                        dict_value_items4 = extract_values_from_data(mat_file_path, dict4)
                        create_mixed_plot(plot_file_path_13, dict_value_items, title='ADC measurements - voltages',
                                          y_names='Voltage')
                        create_mixed_plot(f'{plot_file_path_13[:-4]}_2.png', dict_value_items2,
                                          title='ADC measurements - currents', y_names='Current')
                        create_plot(f'{plot_file_path_13[:-4]}_3.png', dict_value_items3,
                                    'ADC measurements - voltage in over time', mean_sd=True, range=(50, 60))
                        create_plot(f'{plot_file_path_13[:-4]}_4.png', dict_value_items4,
                                    'ADC measurements - current in over time', mean_sd=True, range=(0, 1))
                        plot_paths.append((plot_file_path_13, ''))
                        plot_paths.append((f'{plot_file_path_13[:-4]}_2.png', ''))
                        plot_paths.append((f'{plot_file_path_13[:-4]}_3.png', ''))
                        plot_paths.append((f'{plot_file_path_13[:-4]}_4.png', ''))

                        progress_dialog.setLabelText(f"Building ADC Data Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Vn Histogram"]:
                        plot_file_name = f'{base_name}_vn_histogram.png'
                        plot_file_path = os.path.join(plots_folder, plot_file_name)
                        dict = {'vn': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path, title='Vn')
                        plot_paths.append((plot_file_path, ''))

                        progress_dialog.setLabelText(f"Building Vn Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Servo Histogram"]:
                        plot_file_name_21 = f'{base_name}_servo_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'servo': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='Servo')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Servo Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Bar30 Histogram"]:
                        plot_file_name_21 = f'{base_name}_bar30_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'ba30': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='Bar30')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Bar30 Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["BME Histogram"]:
                        plot_file_name_21 = f'{base_name}_bme_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'bme_back': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='BME')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building BME Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["ADC Histogram"]:
                        plot_file_name_21 = f'{base_name}_adc_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'adc': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='ADC')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building ADC Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Pid pitch Histogram"]:
                        plot_file_name_21 = f'{base_name}_pid_pitch_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'pid': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='Pid pitch')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Pid Pitch Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Pid heading Histogram"]:
                        plot_file_name_21 = f'{base_name}_pid_heading_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'pid': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='Pid heading')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Pid Heading Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Jet Histogram"]:
                        plot_file_name_21 = f'{base_name}_jet_histogram.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'jet': ['timestamp']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_histogram(dict_value_items, plot_file_path_21, title='Jet')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Jet Histogram for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Position & Velocity Uncertainty/Time Plot"]:
                        plot_file_name_21 = f'{base_name}_position_velocity_uncertainty_time_plot.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'vn': ['timestamp', 'posu', 'velu']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_mixed_plot(plot_file_path_21, dict_value_items, title='Pos & Vel Uncertainty Over Time',
                                          y_names='Pos & Vel Uncertainty')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Position & Velocity Uncertainty/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Fix & Num Sats/Time Plot"]:
                        plot_file_name_21 = f'{base_name}_fix_num_sats_time_plot.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'gnssVn200': ['timestamp', 'fix', 'numsats']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        create_mixed_plot(plot_file_path_21, dict_value_items, title='Fix & # of Sats Over Time',
                                          y_names='Fix & # of Sats')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Fix & Num Sats/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Yaw & Status/Time Plot"]:
                        plot_file_name_21 = f'{base_name}_yaw_status_time_plot.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'vn': ['timestamp', 'yaw']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        dict2 = {'status': ['timestamp', 'status_code']}
                        dict_value_items2 = extract_values_from_data(mat_file_path, dict2)
                        create_mixed_plot(plot_file_path_21, dict_value_items, dict_value_items2,
                                          title='Yaw & Status Over Time', y_names='Yaw & Status')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Yaw & Status/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Roll & Status/Time Plot"]:
                        plot_file_name_21 = f'{base_name}_roll_status_time_plot.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'vn': ['timestamp', 'roll']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        dict2 = {'status': ['timestamp', 'status_code']}
                        dict_value_items2 = extract_values_from_data(mat_file_path, dict2)
                        create_mixed_plot(plot_file_path_21, dict_value_items, dict_value_items2,
                                          title='Roll & Status Over Time', y_names='Roll & Status')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Roll & Status/Time Plot for {base_name}...")
                        QApplication.processEvents()

                    if selected_plots[mat_file_path]["Depth & Pitch/Time Plot"]:
                        plot_file_name_21 = f'{base_name}_depth_pitch_time_plot.png'
                        plot_file_path_21 = os.path.join(plots_folder, plot_file_name_21)
                        dict = {'guidance_data_log': ['timestamp', 'pitch']}
                        dict_value_items = extract_values_from_data(mat_file_path, dict)
                        dict2 = {'ba30': ['timestamp', 'depth']}
                        dict_value_items2 = extract_values_from_data(mat_file_path, dict2)
                        create_mixed_plot(plot_file_path_21, dict_value_items, dict_value_items2,
                                          title='Depth & Pitch Over Time', y_names='Depth & Pitch')
                        plot_paths.append((plot_file_path_21, ''))

                        progress_dialog.setLabelText(f"Building Depth & Pitch/Time Plot for {base_name}...")
                        QApplication.processEvents()

                except Exception as e:
                    QMessageBox.critical(None, "Error", f"An error occurred while processing {base_name}: {str(e)}")
                    print(f"Error processing {base_name}: {e}")
                    continue

                # Update progress dialog
                progress_dialog.setValue(idx + 1)

                if progress_dialog.wasCanceled():
                    break

            doc_file_path = os.path.join(output_folder, word_file_name)
            create_word_document(plot_paths, doc_file_path, map_paths)

            # Show completion message
            QMessageBox.information(None, "Success", f"Report generation completed successfully. Files saved to {output_folder}")

            print(f"Processing complete. Files saved to {output_folder}")

        except Exception as e:
            QMessageBox.critical(None, "Error", f"An unexpected error occurred: {str(e)}")
            print(f"An unexpected error occurred: {e}")
        finally:
            progress_dialog.close()

        app.exec_()

    except Exception as e:
        print(f"An error occurred: {e}")

def extract_values_from_data(mat_file_path, dict):
    data = scipy.io.loadmat(mat_file_path)
    pprint(data)
    dict_value_items = {}
    for key, values in dict.items():
        if key not in data:
            print(f"Key '{key}' not found in {mat_file_path}. Skipping...")
            return

        field_data = data[key]
        # Check if the field_data array is empty
        if field_data.size == 0:
            print(f"No data found in table '{key}'. Skipping...")
            return

        for value in values:
            try:
                items = field_data[0][0][value][0]
                dict_value_items[value] = items

            except (IndexError, KeyError):
                print(f"Error accessing data for '{value}' in table '{key}'. Skipping...")
                return

    return dict_value_items

